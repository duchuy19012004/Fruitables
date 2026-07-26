from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\juven\Desktop\Fruitables")
OUT = ROOT / "outputs" / "kich-ban-thuyet-trinh-fruitables-b2c.docx"

GREEN = "5A8A00"
BRIGHT_GREEN = "81C408"
SOFT_GREEN = "F1F8E8"
INK = "1F241B"
MUTED = "5F6B55"
LIGHT = "E8EBE3"


def set_font(run, name="Calibri", size=11, bold=False, color=INK, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_font(run, size=9, color=MUTED)


def configure(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, GREEN, 18, 10),
        ("Heading 2", 13, GREEN, 14, 7),
        ("Heading 3", 12, GREEN, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Script" not in doc.styles:
        style = doc.styles.add_style("Script", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.space_after = Pt(7)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("FRUITABLES B2C | KỊCH BẢN THUYẾT TRÌNH")
    set_font(run, size=9, bold=True, color=GREEN)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Trang ")
    set_font(run, size=9, color=MUTED)
    add_page_number(footer)


def add_label(doc, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(value.upper())
    set_font(run, size=9, bold=True, color=GREEN)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    shade(cell, SOFT_GREEN)
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    set_font(run, size=10, bold=True, color=GREEN)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    run2 = p2.add_run(body)
    set_font(run2, size=10, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_entry(doc, item):
    doc.add_heading(f"Slide {item['n']}: {item['title']}", level=2)
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(7)
    a = meta.add_run(f"Thời lượng gợi ý: {item['time']}  |  ")
    set_font(a, size=10, bold=True, color=GREEN)
    b = meta.add_run(f"Mục đích: {item['purpose']}")
    set_font(b, size=10, color=MUTED)

    add_label(doc, "Lời nói gợi ý")
    p = doc.add_paragraph(style="Script")
    p.paragraph_format.left_indent = Inches(0.12)
    p.add_run(item["script"])

    if item.get("emphasis"):
        add_callout(doc, "Cần nhấn", item["emphasis"])

    if item.get("transition"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(8)
        lead = p.add_run("Chuyển slide: ")
        set_font(lead, size=10, bold=True, color=GREEN)
        text = p.add_run(item["transition"])
        set_font(text, size=10, color=MUTED, italic=True)


slides = [
    {
        "n": 1,
        "title": "Trang bìa",
        "time": "20 giây",
        "purpose": "Chào, giới thiệu đề tài và nhóm thực hiện.",
        "script": "Em xin kính chào thầy cô và các bạn. Nhóm 2 xin trình bày đề tài xây dựng hệ thống thương mại điện tử rau củ quả theo mô hình B2C, với tên hệ thống là Fruitables. Mô hình B2C ở đây có nghĩa là cửa hàng bán trực tiếp cho người tiêu dùng cuối. Trong phần trình bày, nhóm sẽ đi từ bài toán thực tế, các nhóm người dùng, rồi đi sâu vào ba chức năng chính là Chatbox, quản lý giá và tạo combo sản phẩm.",
        "emphasis": "Nói tên đề tài chậm, rõ. Không cần đọc lại toàn bộ tên thành viên nếu đã được giới thiệu trước đó.",
        "transition": "Trước hết, em xin giới thiệu ngắn gọn Fruitables giải quyết việc gì."
    },
    {
        "n": 2,
        "title": "Giới thiệu Fruitables",
        "time": "35 giây",
        "purpose": "Nêu phạm vi của hệ thống.",
        "script": "Fruitables là website bán rau củ quả trực tiếp từ cửa hàng đến khách mua. Với khách hàng, hệ thống hỗ trợ tìm sản phẩm, đặt hàng và theo dõi đơn ở một nơi. Với phía cửa hàng, chủ và nhân viên có khu vực quản trị để xử lý các công việc vận hành. Vì vậy, đây không chỉ là một trang bán hàng, mà là một hệ thống có cả phần mua sắm và phần quản lý phía sau.",
        "emphasis": "Giải thích B2C thật ngắn: cửa hàng bán thẳng cho khách, không qua đại lý.",
        "transition": "Từ phạm vi đó, nhóm xác định một số vấn đề thường gặp khi bán hàng rời rạc."
    },
    {
        "n": 3,
        "title": "Bài toán cần giải quyết",
        "time": "45 giây",
        "purpose": "Làm rõ lý do cần xây dựng hệ thống.",
        "script": "Nếu việc bán rau củ quả được xử lý ở nhiều nơi, khách thường khó tìm đúng loại sản phẩm, mức giá hoặc phân loại phù hợp. Ở phía cửa hàng, giá, tồn kho và trạng thái sản phẩm có thể bị lệch nếu cập nhật không đồng bộ. Đơn hàng, đánh giá và tài khoản khách cũng dễ bị tách rời. Ngoài ra, khi khách đang phân vân trước khi mua, cửa hàng cần có cách hỗ trợ nhanh. Đây là các vấn đề mà Fruitables hướng tới xử lý trong cùng một hệ thống.",
        "emphasis": "Không cần nói đây là vấn đề của mọi cửa hàng. Chỉ nêu đây là tình huống hệ thống được thiết kế để giảm bớt.",
        "transition": "Từ các vấn đề này, nhóm xây dựng một quy trình mua hàng liền mạch hơn."
    },
    {
        "n": 4,
        "title": "Giải pháp tổng thể",
        "time": "45 giây",
        "purpose": "Cho người nghe thấy hành trình B2C đầy đủ.",
        "script": "Quy trình của Fruitables bắt đầu khi khách khám phá sản phẩm. Sau đó khách thêm hàng vào giỏ, đặt hàng và thanh toán, rồi theo dõi việc giao nhận. Khi đơn hoàn tất, hệ thống vẫn có phần chăm sóc sau mua như đánh giá sản phẩm. Trong đồ án này, nhóm chọn trình bày sâu hơn ba phần có logic riêng và tác động trực tiếp đến trải nghiệm mua hàng: Chatbox, quản lý giá và combo sản phẩm.",
        "emphasis": "Chỉ dùng ngón tay đi theo 5 bước trên slide, không đọc lại từng chữ.",
        "transition": "Tiếp theo, em xin giới thiệu ai là người sử dụng các chức năng này."
    },
    {
        "n": 5,
        "title": "Chuyển phần Use case",
        "time": "15 giây",
        "purpose": "Báo người nghe chuyển sang phần vai trò và chức năng.",
        "script": "Phần tiếp theo là use case của hệ thống. Ở phần này, em sẽ chia người dùng thành hai nhóm để dễ theo dõi: nhóm mua hàng và nhóm vận hành cửa hàng.",
        "transition": "Đầu tiên là bức tranh tổng quát của bốn vai trò."
    },
    {
        "n": 6,
        "title": "Use case tổng quát",
        "time": "50 giây",
        "purpose": "Phân nhóm bốn vai trò chính.",
        "script": "Hệ thống có bốn vai trò. Nhóm mua hàng gồm khách vãng lai và khách hàng đã đăng nhập. Khách vãng lai có thể xem, tìm kiếm và thêm sản phẩm vào giỏ. Khi đăng nhập, khách hàng có thêm quyền đặt hàng, thanh toán, theo dõi đơn và đánh giá. Nhóm vận hành gồm nhân viên và chủ cửa hàng. Nhân viên xử lý công việc hằng ngày như sản phẩm, đơn hàng và hỗ trợ khách. Chủ cửa hàng có thêm các quyền điều chỉnh chính sách bán hàng, ví dụ giá, combo và phân quyền.",
        "emphasis": "Giữ ý chính: quyền càng cao thì phạm vi quản lý càng rộng.",
        "transition": "Bây giờ em sẽ nói rõ hơn hành trình của khách mua."
    },
    {
        "n": 7,
        "title": "Use case khách hàng",
        "time": "40 giây",
        "purpose": "Giải thích vì sao khách chưa đăng nhập vẫn có thể bắt đầu mua.",
        "script": "Fruitables cho phép khách bắt đầu mua sắm ngay cả khi chưa đăng nhập. Họ có thể tìm kiếm, lọc sản phẩm, xem đánh giá và thêm hàng vào giỏ. Khi cần hoàn tất giao dịch, khách đăng nhập để đặt hàng và theo dõi trạng thái đơn. Cách làm này giảm bớt bước bắt buộc ở đầu hành trình, vì khách có thể xem trước rồi mới quyết định đăng nhập. Chatbox cũng có thể hỗ trợ ở cả hai trạng thái.",
        "emphasis": "Đừng nói khách vãng lai được thanh toán. Việc đăng nhập được yêu cầu khi cần hoàn tất giao dịch.",
        "transition": "Slide sau sẽ tách riêng các thao tác của khách vãng lai."
    },
    {
        "n": 8,
        "title": "Use case chi tiết: khách vãng lai",
        "time": "35 giây",
        "purpose": "Tóm tắt bốn thao tác trước khi mua.",
        "script": "Với khách vãng lai, có bốn thao tác chính. Thứ nhất là tìm kiếm và lọc để nhanh chóng thu hẹp danh sách sản phẩm. Thứ hai là xem danh sách và chi tiết sản phẩm. Thứ ba là xem đánh giá để có thêm thông tin trước khi chọn mua. Cuối cùng, khách có thể thêm sản phẩm vào giỏ. Mục tiêu của nhóm chức năng này là để khách tìm được thứ mình cần mà chưa phải tạo tài khoản ngay.",
        "emphasis": "Có thể diễn giải lọc theo loại rau, giá hoặc nhu cầu mua. Không cần tự thêm chức năng ngoài slide.",
        "transition": "Sau khi đăng nhập, khách có đầy đủ các chức năng cho cả quá trình mua và sau mua."
    },
    {
        "n": 9,
        "title": "Use case chi tiết: khách hàng",
        "time": "45 giây",
        "purpose": "Nêu các chức năng sau khi khách có tài khoản.",
        "script": "Khách hàng đã đăng nhập sẽ kế thừa các quyền của khách vãng lai và có thêm các chức năng cần cho giao dịch. Họ có thể quản lý hồ sơ, địa chỉ nhận hàng, đặt hàng và thanh toán. Sau khi mua, khách theo dõi trạng thái đơn, có thể hủy đơn khi trạng thái còn cho phép, và đánh giá sản phẩm. Như vậy tài khoản khách hàng giúp hệ thống gắn các hoạt động mua sắm với đúng người dùng và đúng đơn hàng.",
        "emphasis": "Cụm 'khi trạng thái còn cho phép' rất quan trọng, vì không phải đơn nào cũng được hủy.",
        "transition": "Bên cạnh người mua, hệ thống còn có nhóm quản trị để vận hành cửa hàng."
    },
    {
        "n": 10,
        "title": "Use case quản trị",
        "time": "45 giây",
        "purpose": "Phân biệt công việc của nhân viên và chủ cửa hàng.",
        "script": "Ở phía quản trị, nhân viên tập trung vào các việc phát sinh hằng ngày như quản lý sản phẩm, danh mục, đơn hàng, đánh giá, tài khoản khách và báo cáo cơ bản. Chủ cửa hàng có toàn bộ quyền của nhân viên, đồng thời có quyền điều chỉnh chính sách kinh doanh như giá, combo sản phẩm, voucher, giao nhận, SEO và phân quyền. Việc tách vai trò như vậy giúp các thao tác nhạy cảm không bị mở cho tất cả tài khoản quản trị.",
        "emphasis": "Nói rõ: chủ cửa hàng bao gồm quyền nhân viên, không phải một nhóm hoàn toàn tách biệt.",
        "transition": "Hai slide sau sẽ cho thấy use case chi tiết của từng vai trò vận hành."
    },
    {
        "n": 11,
        "title": "Use case nhân viên",
        "time": "30 giây",
        "purpose": "Gói gọn phạm vi công việc thường ngày của nhân viên.",
        "script": "Với nhân viên, nhóm use case được gom thành hai cụm. Cụm thứ nhất là sản phẩm và đơn hàng, tức các công việc trực tiếp phục vụ việc bán. Cụm thứ hai là khách hàng và báo cáo, bao gồm hỗ trợ khách, duyệt hoặc ẩn đánh giá khi cần, và xem thông tin phục vụ vận hành. Điểm cần nhớ là nhân viên làm công việc thường ngày, còn các chính sách lớn vẫn do chủ cửa hàng kiểm soát.",
        "transition": "Tiếp theo là phạm vi rộng hơn của chủ cửa hàng."
    },
    {
        "n": 12,
        "title": "Use case chủ cửa hàng",
        "time": "35 giây",
        "purpose": "Nhấn quyền điều hành và quyền cấu hình.",
        "script": "Chủ cửa hàng có hai nhóm quyền. Nhóm vận hành chung gồm những việc mà nhân viên cũng có thể thực hiện. Nhóm kinh doanh và hệ thống gồm các phần cần kiểm soát cao hơn, như giá, combo, voucher, giao nhận, SEO và phân quyền. Trong phần còn lại của bài trình bày, nhóm đi sâu vào ba chức năng thuộc nhóm này là Chatbox, quản lý giá và combo sản phẩm.",
        "emphasis": "Đây là cầu nối sang ba phần chính, nên nói chậm ở câu cuối.",
        "transition": "Phần đầu tiên là Chatbox hỗ trợ khách hàng."
    },
    {
        "n": 13,
        "title": "Chuyển phần Chatbox",
        "time": "15 giây",
        "purpose": "Giới thiệu mục tiêu của Chatbox.",
        "script": "Ở phần Chatbox, nhóm sẽ bắt đầu từ dữ liệu mà hệ thống lưu, sau đó theo dõi ba chặng xử lý: mở phiên chat, gửi câu hỏi và tạo phản hồi. Hình bên trái chỉ minh họa giao diện mà người dùng nhìn thấy.",
        "transition": "Đầu tiên là cấu trúc dữ liệu của Chatbox."
    },
    {
        "n": 14,
        "title": "Database Chatbox",
        "time": "60 giây",
        "purpose": "Giải thích dữ liệu hội thoại và kho tri thức bằng ngôn ngữ đơn giản.",
        "script": "Database Chatbox được tách thành hai nhóm để dễ quản lý. Nhóm hội thoại gồm Users, ChatSessions và ChatMessages. Hiểu đơn giản, một người dùng có thể có nhiều phiên chat, và một phiên chat có nhiều tin nhắn. Nhóm còn lại là kho tri thức. Sản phẩm và câu hỏi thường gặp được đưa thành các đoạn nội dung gọi là KnowledgeChunks. Khi khách hỏi, hệ thống sẽ tìm các đoạn liên quan trong kho này để hỗ trợ câu trả lời. Trên sơ đồ, PK là khóa chính để nhận diện một bản ghi, còn FK là khóa liên kết giữa các bảng. Người nghe chỉ cần hiểu mối quan hệ chính, không cần đọc từng cột nhỏ.",
        "emphasis": "Nếu bị hỏi 'RAG là gì', có thể trả lời: đó là cách tìm dữ liệu liên quan trước khi AI tạo câu trả lời.",
        "transition": "Từ dữ liệu này, luồng Chatbox bắt đầu khi người dùng mở khung chat."
    },
    {
        "n": 15,
        "title": "Chatbox 1/3: mở chat và khôi phục phiên",
        "time": "50 giây",
        "purpose": "Diễn giải việc tạo hoặc dùng lại phiên chat.",
        "script": "Khi người dùng mở khung chat, giao diện trước hết kiểm tra xem có phiên chat gần nhất hay không. Nếu có, API tải lại phiên đã lưu để người dùng không bị mất mạch hội thoại. Nếu chưa có, hệ thống tạo một phiên mới. Trường hợp không thể khôi phục phiên, giao diện hiển thị lỗi cho người dùng và người dùng có thể đóng chat. Ở đây hệ thống không ép người dùng phải làm tiếp, mà chỉ cho biết tình trạng đang xảy ra.",
        "emphasis": "Theo thứ tự từ trái sang phải của swimlane: người dùng, giao diện, API, rồi phần AI ở phía sau.",
        "transition": "Sau khi có phiên chat, người dùng có thể gửi câu hỏi."
    },
    {
        "n": 16,
        "title": "Chatbox 2/3: gửi câu hỏi và kiểm tra yêu cầu",
        "time": "50 giây",
        "purpose": "Làm rõ bước kiểm tra trước khi xử lý AI.",
        "script": "Người dùng nhập câu hỏi hoặc chọn một gợi ý có sẵn. Ngay lúc đó, giao diện hiển thị tin vừa gửi và báo đang soạn phản hồi để người dùng biết hệ thống đang làm việc. API tiếp tục kiểm tra độ dài nội dung và giới hạn gửi. Nếu yêu cầu không hợp lệ, giao diện trả lỗi ngay. Nếu hợp lệ, hệ thống lưu câu hỏi và chuyển sang phần tìm kiếm tri thức. Bước kiểm tra này giúp Chatbox tránh xử lý những yêu cầu không phù hợp ngay từ đầu.",
        "emphasis": "Không cần nói thuật toán kiểm tra cụ thể. Chỉ cần giải thích đây là bước bảo đảm dữ liệu đầu vào hợp lệ.",
        "transition": "Phần cuối cùng là cách hệ thống tạo phản hồi dựa trên dữ liệu đã có."
    },
    {
        "n": 17,
        "title": "Chatbox 3/3: truy xuất tri thức và phản hồi",
        "time": "60 giây",
        "purpose": "Giải thích RAG và quyền quyết định của người dùng.",
        "script": "Ở chặng này, phần RAG chuyển câu hỏi thành dạng có thể tìm kiếm, rồi tìm và xếp hạng các đoạn kiến thức liên quan. AI dùng các đoạn tìm được để soạn câu trả lời. Nếu dữ liệu chưa đủ, hệ thống thông báo rõ thay vì tự tạo câu trả lời thiếu căn cứ. Cuối cùng, phản hồi được lưu và cập nhật lên màn hình. Người dùng có thể tiếp tục hỏi hoặc đóng chat. Như vậy, Chatbox vừa hỗ trợ nhanh, vừa giữ được ngữ cảnh và lịch sử trao đổi.",
        "emphasis": "Câu quan trọng nhất: AI dựa trên dữ liệu đã tìm được, không trả lời tùy ý.",
        "transition": "Sau Chatbox, nhóm chuyển sang chức năng tác động trực tiếp đến giá bán."
    },
    {
        "n": 18,
        "title": "Chuyển phần Quản lý giá",
        "time": "30 giây",
        "purpose": "Giới thiệu phần giá và lịch áp dụng.",
        "script": "Phần tiếp theo là quản lý giá. Với một cửa hàng bán rau củ quả, giá có thể thay đổi theo từng đợt khuyến mãi hoặc theo thời gian bán hàng. Nếu chỉ sửa trực tiếp một con số, cửa hàng sẽ khó biết khi nào cần khôi phục lại giá cũ và ai là người đã chỉnh sửa. Vì vậy, Fruitables tách việc quản lý giá thành một chức năng riêng. Admin có thể xem giá hiện tại, tạo lịch giảm giá có thời gian bắt đầu và kết thúc, sau đó hệ thống kiểm tra trước khi cho lịch có hiệu lực. Hình bên trái là minh họa giao diện mà admin dùng để theo dõi và lập lịch giá.",
        "emphasis": "Mục tiêu của phần này là quản lý giá có thời hạn và có thể truy vết, không phải chỉ tạo một ô nhập giá mới.",
        "transition": "Để hiểu vì sao có thể theo dõi và khôi phục giá, trước hết em xin giải thích các bảng dữ liệu liên quan."
    },
    {
        "n": 19,
        "title": "Database quản lý giá",
        "time": "1 phút 40 giây",
        "purpose": "Giải thích giá gốc, lịch giá và nhật ký thay đổi.",
        "script": "Ở slide này, em xin chia sơ đồ thành ba nhóm để dễ theo dõi. Nhóm thứ nhất là Users, Products và ProductVariants. Users cho biết ai đang thực hiện thao tác quản trị. Products lưu thông tin sản phẩm gốc, còn ProductVariants dùng khi một sản phẩm có nhiều lựa chọn, ví dụ khác loại hoặc khác quy cách. Mỗi sản phẩm hoặc biến thể có giá ban đầu của nó. Nhóm thứ hai là PriceSchedules. Đây là bảng quan trọng nhất của chức năng này. Một lịch giá sẽ lưu sản phẩm hoặc biến thể được áp dụng, giá trị giảm, thời điểm bắt đầu, thời điểm kết thúc và người tạo lịch. Vì lịch được lưu riêng, admin có thể tạo trước một chương trình giảm giá thay vì phải chờ đến đúng giờ rồi chỉnh tay. Nhóm cuối cùng là ProductLogs. Bảng này lưu lại hành động thay đổi để khi cần kiểm tra, cửa hàng biết sản phẩm nào đã thay đổi và ai là người thao tác. Từ ba nhóm đó, giá có hiệu lực mà khách nhìn thấy được hiểu là giá gốc kết hợp với lịch giảm đang chạy. Khi lịch hết hạn, hệ thống quay về giá gốc. Giá ban đầu vẫn còn trong dữ liệu, nên không bị mất sau một đợt khuyến mãi.",
        "emphasis": "Không cần đọc từng cột. Hãy chỉ lần lượt vào ba nhóm: sản phẩm và biến thể, lịch giá, nhật ký thay đổi. Nếu bị hỏi PK và FK, trả lời PK là mã nhận diện của bản ghi, FK là mã dùng để liên kết các bảng.",
        "transition": "Sau khi đã biết dữ liệu được lưu ở đâu, em xin đi vào cách admin tạo một lịch giảm giá."
    },
    {
        "n": 20,
        "title": "Quản lý giá 1/3: tạo lịch giảm",
        "time": "1 phút 10 giây",
        "purpose": "Mô tả thao tác khởi đầu của admin.",
        "script": "Luồng bắt đầu khi admin mở trang quản lý giá. Admin chọn sản phẩm hoặc chọn một biến thể cụ thể cần áp dụng. Bước này quan trọng vì một sản phẩm có thể có nhiều biến thể, và không phải lúc nào cũng giảm giá cho toàn bộ biến thể. Sau khi admin chọn, hệ thống hiển thị giá hiện tại cùng những lịch giảm giá đã tồn tại. Mục đích không phải chỉ để xem lại, mà để admin biết sản phẩm đang có lịch nào, thời gian nào đã được đặt và có cần tạo lịch mới hay không. Tiếp theo, admin nhập mức giảm và khoảng thời gian áp dụng, gồm thời điểm bắt đầu và kết thúc. Khi đã kiểm tra thông tin trên form, admin bấm xác nhận tạo lịch. Ở thời điểm này, hệ thống mới nhận dữ liệu để kiểm tra. Giá chưa thay đổi ngay chỉ vì admin vừa bấm tạo lịch.",
        "emphasis": "Dùng từ admin vì activity diagram dùng vai trò này. Nhấn rõ ba việc: chọn đúng đối tượng áp dụng, xem giá và lịch hiện tại, rồi mới nhập thời gian.",
        "transition": "Sau khi admin gửi thông tin, hệ thống cần kiểm tra lịch đó có phù hợp hay không trước khi lưu."
    },
    {
        "n": 21,
        "title": "Quản lý giá 2/3: kiểm tra và chờ áp dụng",
        "time": "1 phút 20 giây",
        "purpose": "Giải thích nhánh lỗi và trạng thái chờ.",
        "script": "Sau khi nhận dữ liệu, hệ thống kiểm tra ba điểm chính. Thứ nhất là mức giảm có hợp lý theo quy tắc của cửa hàng hay không. Thứ hai là khoảng thời gian có đầy đủ và có thứ tự đúng không, ví dụ thời điểm kết thúc không thể nằm trước thời điểm bắt đầu. Thứ ba là lịch mới có bị chồng lên một lịch khác của cùng sản phẩm hoặc biến thể hay không. Nếu có vấn đề, hệ thống trả thông báo lỗi về form. Admin nhìn thấy lỗi và có quyền tự quyết định: chỉnh lại thông tin rồi lưu lại, hoặc dừng tác vụ. Hệ thống không tự sửa dữ liệu thay admin. Nếu mọi thông tin hợp lệ, lịch được lưu ở trạng thái chờ. Có thể hiểu trạng thái chờ là lịch đã tồn tại trong database, nhưng chưa tác động đến giá mà khách đang nhìn thấy. Việc lưu và việc áp dụng là hai mốc khác nhau. Điều này giúp admin có thể chuẩn bị chương trình giảm giá từ trước.",
        "emphasis": "Nhấn rõ hai nhánh: lỗi thì trả về form để admin quyết định, hợp lệ thì lưu chờ. Câu 'đã lưu nhưng chưa áp dụng' cần nói chậm.",
        "transition": "Khi đến đúng thời điểm đã đặt, hệ thống mới bắt đầu áp dụng lịch này."
    },
    {
        "n": 22,
        "title": "Quản lý giá 3/3: tự động áp dụng và kết thúc",
        "time": "1 phút 15 giây",
        "purpose": "Mô tả cơ chế áp dụng và khôi phục giá gốc.",
        "script": "Ở chặng cuối, hệ thống định kỳ kiểm tra các lịch đang chờ và so sánh với thời điểm hiện tại. Khi lịch đến giờ bắt đầu, mức giảm của lịch được đưa vào giá có hiệu lực. Lúc này khách vào trang sản phẩm sẽ nhìn thấy giá đã được cập nhật theo chương trình giảm. Admin không cần quay lại đúng thời điểm đó để bật giá thủ công. Hệ thống tiếp tục theo dõi đến khi lịch hết hạn. Khi thời điểm kết thúc đến, hệ thống dừng áp dụng mức giảm, đưa giá hiển thị trở về giá gốc và đánh dấu lịch là đã hoàn thành. Điểm cần lưu ý là giá gốc không bị ghi đè. Hệ thống chỉ thay đổi giá có hiệu lực trong khoảng thời gian của lịch. Nhờ vậy, sau mỗi đợt giảm giá, cửa hàng vẫn có thể trở về mức giá ban đầu một cách rõ ràng và có lịch sử để kiểm tra lại.",
        "emphasis": "Không khẳng định công nghệ chạy định kỳ cụ thể. Chỉ nói hệ thống tự kiểm tra theo lịch. Có thể chốt phần này bằng câu: giá gốc được giữ lại, lịch giá chỉ có hiệu lực trong đúng khoảng thời gian đã đặt.",
        "transition": "Phần tiếp theo là tính năng tạo combo sản phẩm, nơi giá hiện hành của sản phẩm cũng được sử dụng để tính tổng giá combo."
    },
    {
        "n": 23,
        "title": "Chuyển phần Combo sản phẩm",
        "time": "30 giây",
        "purpose": "Giới thiệu mục tiêu của combo.",
        "script": "Tiếp theo là chức năng tạo combo sản phẩm. Có thể hiểu combo là một gói bán gồm nhiều sản phẩm được chọn cùng nhau. Ví dụ, cửa hàng có thể tạo một combo rau củ dùng cho bữa ăn, trong đó có nhiều loại sản phẩm và mỗi loại có số lượng riêng. Phần này không chỉ nói về việc gom sản phẩm vào một gói. Nhóm cũng trình bày cách hệ thống lấy giá hiện tại của từng sản phẩm, cách admin nhập thông tin, và cách giao diện phản hồi trong trường hợp dữ liệu sai hoặc tạo combo thành công. Hình bên trái là minh họa giao diện chọn sản phẩm để đưa vào combo.",
        "emphasis": "Nói rõ combo là một gói bán, không phải là một sản phẩm mới tách hoàn toàn khỏi các sản phẩm bên trong.",
        "transition": "Trước khi đi vào thao tác tạo, em xin giải thích các bảng dữ liệu dùng cho combo."
    },
    {
        "n": 24,
        "title": "Database combo sản phẩm",
        "time": "1 phút 35 giây",
        "purpose": "Giải thích bảng Combos, ComboItems và cách tính giá.",
        "script": "Sơ đồ này có bốn bảng cần chú ý. Bảng Combos lưu phần thông tin chung của gói bán, như tên combo, mô tả, đường dẫn, trạng thái hoạt động. Đây là phần mà khách sẽ nhìn thấy khi cửa hàng công bố combo. Bảng Products lưu từng sản phẩm gốc, còn ProductVariants dùng khi một sản phẩm có các lựa chọn cụ thể hơn, ví dụ một biến thể khác hoặc một đơn vị bán khác. Ở giữa mối quan hệ này là bảng ComboItems. Có thể hiểu ComboItems là danh sách các món nằm trong một combo. Mỗi dòng trong bảng này cho biết dòng đó thuộc combo nào, chọn sản phẩm nào, có chọn biến thể hay không và số lượng là bao nhiêu. Vì thế, một combo có thể có nhiều ComboItems. Khi cần tính tổng giá, hệ thống lấy giá đang có hiệu lực của từng sản phẩm hoặc biến thể rồi nhân với số lượng. Cách tính này giúp combo phản ánh giá hiện tại thay vì giữ một con số cũ khi giá sản phẩm đã thay đổi.",
        "emphasis": "Giải thích ComboItems là bảng trung gian. Đây là ý dễ hiểu nhất cho người ngoài ngành. Có thể nói: bảng Combos là thông tin của gói, còn ComboItems là danh sách sản phẩm bên trong gói.",
        "transition": "Sau khi đã có cấu trúc dữ liệu, em sẽ đi vào chặng đầu tiên khi admin mở chức năng combo."
    },
    {
        "n": 25,
        "title": "Combo 1/3: mở quản lý và chuẩn bị form",
        "time": "1 phút 10 giây",
        "purpose": "Mô tả màn hình đầu tiên khi admin vào chức năng.",
        "script": "Chặng đầu tiên bắt đầu khi admin chọn menu combo sản phẩm trong khu vực quản trị. Giao diện mở màn hình quản lý combo để admin nhìn thấy danh sách các combo đang có. Ở phía hệ thống, dữ liệu được lấy về gồm danh sách combo và tổng giá hiện tại của từng combo. Tổng giá này có thể thay đổi nếu giá của sản phẩm bên trong combo thay đổi, nên hệ thống cần tính lại theo dữ liệu hiện hành. Khi admin bấm nút thêm combo, giao diện chuyển sang form tạo mới. Trước khi form hiển thị đầy đủ, hệ thống lấy danh sách sản phẩm và biến thể đang được bán. Danh sách đó được đưa vào các ô chọn để admin chọn đúng sản phẩm, đúng biến thể và số lượng. Cách chuẩn bị này giúp form không hiển thị các sản phẩm đã ngừng bán hoặc không còn phù hợp để đưa vào combo.",
        "emphasis": "Ở slide này, nhìn từ trái sang phải để thấy admin, giao diện và hệ thống trao đổi với nhau. Nói rõ hệ thống chuẩn bị danh sách trước khi admin nhập dữ liệu.",
        "transition": "Khi form đã sẵn sàng, admin bắt đầu nhập thông tin của combo mới."
    },
    {
        "n": 26,
        "title": "Combo 2/3: nhập thông tin và kiểm tra",
        "time": "1 phút 15 giây",
        "purpose": "Diễn giải dữ liệu admin nhập và bước gửi kiểm tra.",
        "script": "Ở chặng thứ hai, form hiển thị hai nhóm thông tin. Nhóm đầu là thông tin chung của combo, gồm tên, mô tả, ảnh và trạng thái hoạt động. Nhóm thứ hai là danh sách sản phẩm bên trong combo. Với từng dòng, admin chọn sản phẩm, có thể chọn biến thể nếu cần, rồi nhập số lượng. Ví dụ, nếu một combo có hai loại rau và một loại củ, admin sẽ tạo các mục riêng trong danh sách chi tiết thay vì chỉ nhập một con số tổng. Khi đã hoàn tất, admin bấm nút lưu. Giao diện gửi toàn bộ dữ liệu combo lên hệ thống. Điều cần phân biệt ở đây là giao diện chịu trách nhiệm cho admin nhập và nhìn lại dữ liệu, còn hệ thống mới là nơi kiểm tra dữ liệu có hợp lệ hay không. Nhờ tách hai vai trò đó, mọi kết quả kiểm tra đều được trả về rõ ràng trên cùng form mà admin đang sử dụng.",
        "emphasis": "Không cần nêu một quy tắc kiểm tra cụ thể nếu slide không có. Hãy giữ trọng tâm là dữ liệu được nhập trên form rồi gửi lên hệ thống để xác nhận.",
        "transition": "Sau khi kiểm tra, luồng đi theo một trong hai nhánh: báo lỗi trên form hoặc lưu combo thành công."
    },
    {
        "n": 27,
        "title": "Combo 3/3: xử lý lỗi, lưu thành công và tạo tiếp",
        "time": "1 phút 35 giây",
        "purpose": "Làm rõ trải nghiệm khi lỗi và quyền lựa chọn sau khi thành công.",
        "script": "Đây là chặng cuối và cũng là chặng có hai nhánh rõ nhất. Nếu dữ liệu không hợp lệ, hệ thống không xóa form và cũng không tự chỉnh sửa thay admin. Hệ thống trả lỗi ngay tại vị trí cần chú ý trên form, đồng thời giữ lại tên combo, mô tả, sản phẩm và số lượng mà admin đã nhập. Admin có thể xem lỗi, chỉnh sửa rồi bấm lưu lại. Nếu admin không muốn tiếp tục, họ có thể dừng tác vụ. Nhánh còn lại là khi dữ liệu hợp lệ. Hệ thống lưu thông tin combo cùng danh sách ComboItems, sau đó giao diện hiển thị thông báo thành công và cập nhật bảng danh sách combo. Admin có thể xem combo vừa tạo để kiểm tra kết quả. Sau đó admin có quyền chọn tạo tiếp một combo khác hoặc kết thúc. Điểm nhóm muốn làm rõ là hệ thống không ép người dùng phải tiếp tục sau khi thành công, cũng không ép người dùng phải sửa ngay khi gặp lỗi. Hệ thống hiển thị kết quả, còn quyết định tiếp theo thuộc về admin.",
        "emphasis": "Đây là slide nên nói chậm nhất trong phần combo. Nhấn hai ý: lỗi hiển thị ngay trên form và giữ lại dữ liệu; thành công thì admin được chọn tạo tiếp hoặc kết thúc.",
        "transition": "Như vậy, phần combo đã hoàn tất. Em xin chuyển sang phần kết luận để tổng hợp giá trị của toàn hệ thống."
    },
    {
        "n": 28,
        "title": "Chuyển phần Kết luận",
        "time": "30 giây",
        "purpose": "Khép lại phần chức năng và chuyển sang giá trị của hệ thống.",
        "script": "Qua ba phần vừa trình bày, Chatbox, quản lý giá và combo sản phẩm, có thể thấy Fruitables không chỉ có phần giao diện để khách mua hàng. Mỗi chức năng đều có dữ liệu riêng, có bước kiểm tra và có luồng xử lý ở phía quản trị. Phần cuối cùng, nhóm sẽ nhìn lại giá trị của hệ thống ở hai phía: khách mua nhận được gì, và cửa hàng vận hành thuận lợi hơn ở điểm nào.",
        "transition": "Đây là phần tổng kết của nhóm."
    },
    {
        "n": 29,
        "title": "Kết luận",
        "time": "1 phút 20 giây",
        "purpose": "Tóm tắt giá trị của ba chức năng theo hai góc nhìn.",
        "script": "Ở phía khách hàng, hệ thống hỗ trợ toàn bộ hành trình mua rau củ quả từ lúc tìm kiếm đến sau khi nhận hàng. Chatbox giúp khách hỏi nhanh khi còn phân vân về sản phẩm hoặc thông tin mua. Quản lý giá giúp khách thấy mức giá đúng theo thời điểm, đặc biệt khi cửa hàng có lịch giảm giá. Combo giúp khách chọn nhiều sản phẩm theo một gói thay vì phải tìm từng sản phẩm riêng lẻ. Ở phía cửa hàng, các dữ liệu này được quản lý tập trung. Chủ và nhân viên có thể xử lý sản phẩm, đơn hàng và thông tin khách, đồng thời có thể kiểm tra lại các thay đổi quan trọng như lịch giá. Các activity diagram cũng cho thấy hệ thống luôn có bước kiểm tra trước khi lưu. Khi có lỗi, giao diện trả thông báo rõ ràng và giữ quyền quyết định cho người quản trị. Vì vậy, Fruitables kết nối được trải nghiệm mua hàng phía trước với công việc vận hành phía sau của cửa hàng.",
        "emphasis": "Kết luận theo hai phía: khách mua và cửa hàng. Không mở thêm tính năng mới ở đoạn này.",
        "transition": "Em xin cảm ơn thầy cô và các bạn đã lắng nghe."
    },
    {
        "n": 30,
        "title": "Cảm ơn và hỏi đáp",
        "time": "20 giây",
        "purpose": "Kết thúc lịch sự và mời câu hỏi.",
        "script": "Phần trình bày của nhóm đến đây là hết. Nhóm 2 xin cảm ơn thầy cô và các bạn đã lắng nghe. Nhóm sẵn sàng trả lời các câu hỏi liên quan đến use case, database hoặc ba chức năng chính đã trình bày.",
        "emphasis": "Dừng lại, nhìn về phía người nghe và chờ câu hỏi. Không nói thêm nếu chưa có câu hỏi."
    },
]


def build_document(items, filename, title, subtitle):
    doc = Document()
    configure(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(42)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(title)
    set_font(r, size=28, bold=True, color=GREEN)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Fruitables - Hệ thống thương mại điện tử rau củ quả B2C")
    set_font(r, size=17, bold=True, color=INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(subtitle)
    set_font(r, size=12, color=MUTED)
    add_callout(doc, "Cách dùng", "Mỗi slide có một lời nói gợi ý. Người trình bày có thể đọc gần như nguyên văn trong lần tập đầu, sau đó nói theo ý để tự nhiên hơn. Phần 'Cần nhấn' giúp tránh giải thích quá sâu hoặc nói sai trọng tâm.")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("Thời lượng dự kiến: 15 đến 20 phút, chưa tính phần hỏi đáp.")
    set_font(r, size=11, bold=True, color=GREEN)
    doc.add_page_break()

    doc.add_heading("Mở đầu", level=1)
    add_callout(doc, "Gợi ý trước khi trình bày", "Không cần học thuộc từng chữ. Hãy đọc trước kịch bản một lần, nắm các từ in đậm ở phần 'Cần nhấn', và nhìn vào sơ đồ theo hướng từ trái sang phải hoặc từ trên xuống dưới. Khi gặp PK và FK trong database, chỉ cần nói PK là khóa chính và FK là khóa liên kết.")

    groups = {
        1: "Phần 1: Giới thiệu đề tài",
        5: "Phần 2: Use case hệ thống",
        13: "Phần 3: Chatbox hỗ trợ khách hàng",
        18: "Phần 4: Quản lý giá",
        23: "Phần 5: Combo sản phẩm",
        28: "Phần 6: Kết luận",
    }
    for item in items:
        if item["n"] in groups:
            doc.add_heading(groups[item["n"]], level=1)
        add_entry(doc, item)

    output = OUT.parent / filename
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(output)


def main():
    build_document(
        slides,
        "kich-ban-thuyet-trinh-fruitables-b2c.docx",
        "KỊCH BẢN THUYẾT TRÌNH",
        "Dành cho người thuyết trình chưa nắm rõ dự án",
    )
    build_document(
        [item for item in slides if item["n"] <= 17],
        "kich-ban-1-mo-dau-den-het-chatbox.docx",
        "KỊCH BẢN 1",
        "Từ mở đầu đến hết phần Chatbox, slide 1 đến slide 17",
    )
    build_document(
        [item for item in slides if 18 <= item["n"] <= 22],
        "kich-ban-2-phan-quan-ly-gia.docx",
        "KỊCH BẢN 2",
        "Phần Quản lý giá, slide 18 đến slide 22",
    )
    build_document(
        [item for item in slides if item["n"] >= 23],
        "kich-ban-3-combo-den-ket-thuc.docx",
        "KỊCH BẢN 3",
        "Từ phần Combo sản phẩm đến hết bài, slide 23 đến slide 30",
    )


if __name__ == "__main__":
    main()
